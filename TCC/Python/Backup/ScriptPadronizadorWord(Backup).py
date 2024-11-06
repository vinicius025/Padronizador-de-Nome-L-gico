import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from SciprtPadronizadorAuxiliar import ler_abreviacoes, padronizar_nome, formatar_tabela, adicionar_estilo_titulo, adicionar_quebra_pagina
from datetime import datetime  # Biblioteca para obter a data atual

# Função para adicionar bordas a uma célula
def adicionar_bordas(celula):
    tc_pr = celula._element.get_or_add_tcPr()
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tc_pr.append(border)

# Função principal para aplicar a padronização
def aplicar_padronizacao(caminho_abreviacoes, caminho_planilha, caminho_saida):
    # Ler as abreviações
    abreviacoes = ler_abreviacoes(caminho_abreviacoes)

    # Ler a planilha Excel
    df = pd.read_excel(caminho_planilha)

    # Substituir 'nan' por 'Não informado'
    df = df.fillna('Não informado')

    # Obter a data atual
    data_atual = datetime.now().strftime("%d/%m/%Y")

    # Criar um documento Word
    doc = Document()

    # Adicionar o cabeçalho com as informações
    tabela = doc.add_table(rows=7, cols=2) 
    tabela.style = 'Table Grid'

    # Célula para a logo
    celula_logo = tabela.cell(0, 0)
    paragrafo_logo = celula_logo.paragraphs[0]
    paragrafo_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_logo = paragrafo_logo.add_run('Logo da Empresa')
    run_logo.font.size = Pt(16)
    adicionar_bordas(celula_logo)

    # Texto ao lado da logo
    celula_texto = tabela.cell(0, 1)
    paragrafo_texto = celula_texto.paragraphs[0]
    paragrafo_texto.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run_texto = paragrafo_texto.add_run('Dicionario de Dados')
    run_texto.bold = True
    run_texto.font.size = Pt(16)

    # Células com informações
    celulas = [
        ("Nome da Empresa:", ""),
        ("Nome do Sistema:", ""),
        ("Analista de Dados Responsável:", ""),
        ("Analista de Projeto:", ""),
        ("Vice Presidencia:", ""),
        ("Data da Padronização:", data_atual),
    ]

    for i, (titulo, espaco) in enumerate(celulas, start=1):
        celula_titulo = tabela.cell(i, 0)
        celula_espaco = tabela.cell(i, 1)
        paragrafo_titulo = celula_titulo.paragraphs[0]
        run_titulo = paragrafo_titulo.add_run(titulo)
        run_titulo.bold = True
        run_titulo.font.size = Pt(12)
        paragrafo_espaco = celula_espaco.paragraphs[0]
        run_espaco = paragrafo_espaco.add_run(espaco)
        run_espaco.font.size = Pt(12)

    # Adicionar quebra de página para que o conteúdo fique na segunda página
    doc.add_page_break()

    # Adicionar o estilo de título
    adicionar_estilo_titulo(doc, 'Título 1')

    # Adicionar o cabeçalho principal ao documento
    titulo_principal = doc.add_heading('Resultados da Padronização', level=1)
    titulo_principal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Variável para controlar o número de tabelas por página
    tabelas_por_pagina = 2
    contador_tabelas = 0

    # Adicionar informações de cada linha da planilha ao documento
    for index, row in df.iterrows():
        # Obter o nome da tabela da coluna 'TABLE BUSINESS NAME' para exibir ao lado de "Linha X"
        nome_tabela = row['TABLE BUSINESS NAME'] if 'TABLE BUSINESS NAME' in df.columns else "Sem Nome"
        
        # Adicionar uma quebra de página após duas tabelas
        if contador_tabelas >= tabelas_por_pagina:
            doc.add_page_break()
            contador_tabelas = 0  # Reiniciar o contador

        # Definir o título como "Linha X - NOME_TABELA"
        doc.add_heading(f"Linha {index + 1} - {nome_tabela}", level=2)
        
        table = doc.add_table(rows=0, cols=2)  # Criar tabela com 2 colunas (remover "Nome Padronizado")
        formatar_tabela(table)  # Formatar a tabela

        # Adicionar cabeçalhos das colunas
        hdr_cells = table.add_row().cells
        run = hdr_cells[0].paragraphs[0].add_run('Campo')
        run.bold = True
        run.font.size = Pt(10)
        hdr_cells[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        hdr_cells[1].text = 'Valor'
        hdr_cells[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for col in df.columns:
            row_cells = table.add_row().cells
            run = row_cells[0].paragraphs[0].add_run(col)
            run.bold = True
            run.font.size = Pt(10)
            run = row_cells[1].paragraphs[0].add_run(str(row[col]))
            run.font.size = Pt(10)

            # Adicionar nova linha "Nome_Padronizado" abaixo de "COLUMN BUSINESS NAME"
            if col == 'COLUMN BUSINESS NAME':
                nome_padronizado = padronizar_nome(str(row[col]), abreviacoes).upper()
                nome_padronizado_row = table.add_row().cells
                run = nome_padronizado_row[0].paragraphs[0].add_run("Nome_Padronizado")
                run.bold = True
                run.font.size = Pt(10)
                nome_padronizado_row[0].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento padrão à esquerda
                run = nome_padronizado_row[1].paragraphs[0].add_run(nome_padronizado)
                run.font.size = Pt(10)
                nome_padronizado_row[1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento padrão à esquerda

        doc.add_paragraph('---')
        
        # Incrementar o contador de tabelas
        contador_tabelas += 1

    # Salvar o documento Word
    doc.save(caminho_saida)
    print("Padronização concluída e salva em um arquivo Word.")
