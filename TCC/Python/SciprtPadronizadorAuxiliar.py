
import unicodedata  # Biblioteca para remover acentos
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Função para carregar as abreviações de um arquivo de texto
def ler_abreviacoes(caminho_abreviacoes):
    abreviacoes = {}
    with open(caminho_abreviacoes, 'r') as file:
        for linha in file:
            if ' = ' in linha:  # Verifica se a linha está no formato esperado
                palavra, abreviacao = linha.strip().split(' = ')
                # Normalize the key (remove accents, convert to uppercase) for consistent matching
                palavra_normalizada = remover_acentos(palavra).upper()
                abreviacoes[palavra_normalizada] = abreviacao.strip()
    return abreviacoes

# Função para remover acentos
def remover_acentos(texto):
    nfkd = unicodedata.normalize('NFKD', texto)
    texto_sem_acento = ''.join([c for c in nfkd if not unicodedata.combining(c)])
    return texto_sem_acento

# Função para aplicar as regras de padronização a múltiplas palavras
def padronizar_nome(nome, abreviacoes):
    conectores = {'do', 'da', 'de', 'com', 'ou'}
    
    # Remove conectores e normaliza cada palavra (remove acentos e converte para maiúsculas)
    palavras = [
        remover_acentos(palavra).upper() for palavra in nome.split() if palavra.lower() not in conectores
    ]
    
    # Aplica as abreviações usando a forma normalizada
    palavras_abreviadas = [abreviacoes.get(palavra, palavra) for palavra in palavras]
    nome_padronizado = '_'.join(palavras_abreviadas)
    
    return nome_padronizado

def verificar_nomes(caminho_abreviacoes, lista_nomes):
    abreviacoes = ler_abreviacoes(caminho_abreviacoes)
    nenhum_nome_maior = True
    resultado_verificacao = ""

    for nome in lista_nomes:
        nome_padronizado = padronizar_nome(nome, abreviacoes)
        if len(nome_padronizado) > 30:
            resultado_verificacao += f"Aviso: O nome padronizado '{nome_padronizado}' ultrapassa 30 caracteres ({len(nome_padronizado)} caracteres).\n"
            nenhum_nome_maior = False

    if nenhum_nome_maior:
        resultado_verificacao += "Nenhum nome padronizado ultrapassa 30 caracteres.\n"
    
    return resultado_verificacao


# Função para formatar a tabela
def formatar_tabela(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = 300000  # Ajusta a largura de cada célula
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # Define o tamanho da fonte como 10
            # Adicionar borda a cada célula
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')  # Tamanho da borda
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)
            tcPr.append(tcBorders)

# Adicionar estilo de título ao documento Word
def adicionar_estilo_titulo(doc, estilo, nome_fonte='Arial', tamanho_fonte=24, negrito=True):
    estilos = doc.styles
    estilo_titulo = estilos.add_style(estilo, 1)
    fonte = estilo_titulo.font
    fonte.name = nome_fonte
    fonte.size = Pt(tamanho_fonte)
    fonte.bold = negrito

# Função para adicionar quebra de página
def adicionar_quebra_pagina(doc):
    doc.add_paragraph().paragraph_format.page_break_before = True


from docx import Document

# Função para validar nomes padronizados no arquivo .docx
def verificar_nomes_docx(caminho_abreviacoes, caminho_docx):
    # Carregar as abreviações
    abreviacoes = ler_abreviacoes(caminho_abreviacoes)
    nenhum_nome_maior = True  # Flag para verificar se há algum nome maior que 30 caracteres
    
    # Carregar o arquivo .docx
    doc = Document(caminho_docx)
    
    # Iterar pelas tabelas no documento para encontrar entradas "Nome_Padronizado"
    for table in doc.tables:
        for row in table.rows:
            # Verifica se a linha contém "Nome_Padronizado" e valida seu comprimento
            if "Nome_Padronizado" in row.cells[0].text:
                nome_padronizado = row.cells[1].text.strip()  # Obtém o nome associado
                if len(nome_padronizado) > 30:
                    print(f"Aviso: O nome padronizado '{nome_padronizado}' ultrapassa 30 caracteres ({len(nome_padronizado)} caracteres).")
                    nenhum_nome_maior = False

    if nenhum_nome_maior:
        print("Nenhum nome padronizado ultrapassa 30 caracteres.")
